import streamlit as st
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import pandas as pd
import io
import requests

st.title("🧬 Generátor genetické zprávy")

# --- Načtení XLSX z GitHubu ---
url = "https://github.com/matej-bartos/gen/raw/main/Varianty.xlsx"
try:
    response = requests.get(url)
    response.raise_for_status()
    df_all = pd.read_excel(io.BytesIO(response.content), sheet_name="List1")
except Exception as e:
    st.error(f"❌ Chyba při načítání Excelu z GitHubu: {e}")
    st.stop()

# --- Validace sloupců ---
required_cols = ["Sekce", "Gen", "Genotyp", "Interpretace"]
if list(df_all.columns[:4]) != required_cols:
    st.error(f"❌ Soubor musí obsahovat sloupce: {', '.join(required_cols)}")
    st.stop()

df_all = df_all.dropna(subset=required_cols)

# --- Výběr genů podle sekcí ---
vybrane = {}
for sekce in df_all["Sekce"].unique():
    st.subheader(sekce)
    df_sekce = df_all[df_all["Sekce"] == sekce]
    for gen in sorted(set(df_sekce["Gen"])):
        moznosti = df_sekce[df_sekce["Gen"] == gen]["Genotyp"].dropna().astype(str).unique().tolist()
        if moznosti:
            zvolene = st.multiselect(f"{gen}", moznosti, key=gen)
            if zvolene:
                vybrane[gen] = zvolene

# --- Pomocná funkce: pozadí buňky ---
def set_cell_background(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

# --- Generování zprávy ---
if vybrane:
    vysledky = []
    for gen, seznam in vybrane.items():
        for g in seznam:
            z = df_all[(df_all["Gen"] == gen) & (df_all["Genotyp"] == g)].iloc[0]
            vysledky.append(z)
    df_final = pd.DataFrame(vysledky)

    try:
        doc = Document("Vysledkova_zprava.docx")
    except Exception as e:
        st.error(f"❌ Nepodařilo se načíst šablonu: {e}")
        st.stop()

    insert_index = None
    for i, para in enumerate(doc.paragraphs):
        if "TABULKA" in para.text:
            insert_index = i
            doc.paragraphs[i].text = ""
            break
    if insert_index is None:
        st.error("❌ Text 'TABULKA' nebyl nalezen v šabloně.")
        st.stop()

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.autofit = True

    headers = ["GEN", "VÝSLEDNÁ VARIANTA", "INTERPRETACE"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.size = Pt(10)

    # --- Zachování pořadí genů dle Excelu ---
    poradi_genu = df_all["Gen"].drop_duplicates().tolist()

    for sekce in df_final["Sekce"].unique():
        df_sekce = df_final[df_final["Sekce"] == sekce]

        # Nadpis sekce
        row = table.add_row()
        merged = row.cells[0].merge(row.cells[1]).merge(row.cells[2])
        merged.text = sekce
        set_cell_background(merged, "00FFFF")
        para = merged.paragraphs[0]
        run = para.runs[0]
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 32, 96)
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        for gen in [g for g in poradi_genu if g in df_sekce["Gen"].values]:
            df_gen = df_sekce[df_sekce["Gen"] == gen]
            first_row_idx = len(table.rows)

            for _, row_data in df_gen.iterrows():
                row_cells = table.add_row().cells
                row_cells[1].text = str(row_data["Genotyp"])
                row_cells[2].text = str(row_data["Interpretace"])

                # Genotyp = Arial 10 pt
                for run in row_cells[1].paragraphs[0].runs:
                    run.font.name = 'Arial'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
                    run.font.size = Pt(10)

                # Interpretace = Arial 9 pt, tučně, modře
                for run in row_cells[2].paragraphs[0].runs:
                    run.font.name = 'Arial'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
                    run.font.size = Pt(9)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0, 32, 96)

            # Sloučení buňky GEN
            last_row_idx = len(table.rows) - 1
            if last_row_idx > first_row_idx:
                cell_to_merge = table.rows[first_row_idx].cells[0]
                for r in range(first_row_idx + 1, last_row_idx + 1):
                    cell_to_merge.merge(table.rows[r].cells[0])
                cell_to_merge.text = gen
                for para in cell_to_merge.paragraphs:
                    para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    for run in para.runs:
                        run.font.name = 'Arial'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
                        run.font.size = Pt(10)
            else:
                table.rows[first_row_idx].cells[0].text = gen
                for run in table.rows[first_row_idx].cells[0].paragraphs[0].runs:
                    run.font.name = 'Arial'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
                    run.font.size = Pt(10)

    tbl = table._element
    doc.paragraphs[insert_index]._element.addnext(tbl)

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)

    st.download_button(
        label="📄 Stáhnout hotovou zprávu",
        data=output,
        file_name="Geneticka_zprava.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
else:
    st.info("✅ Vyber alespoň jeden genotyp pro generování zprávy.")
